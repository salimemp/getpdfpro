"""Local PDF → Office conversion adapter (always-available fallback).

Uses PyMuPDF for text/structure extraction and python-docx for the
.docx writer. Quality is best-effort: text + basic structure
(headings, simple tables) only. Multi-column layouts, math, and
complex tables will NOT round-trip faithfully.

This is the same code that powers the current /to-word endpoint. It
was extracted into an adapter so the router can use a 3-tier cascade
without duplicating logic.

Honest accuracy: 70-80% for clean PDFs. For higher quality, the cascade
falls through to LibreOffice (90-95%) or Adobe (95-99%).
"""

from __future__ import annotations

import io
import time
from typing import Any

import fitz  # PyMuPDF

from . import AdapterResult, ConversionAdapter, ConversionError, OutputFormat


class LocalAdapter:
    name = "local"
    description = "PyMuPDF + python-docx, runs in-process. Best-effort, always free."
    quality_score = 70

    async def is_available(self) -> bool:
        # Local is always available — no external dependency.
        return True

    async def convert(
        self,
        pdf_bytes: bytes,
        output_format: OutputFormat = "docx",
        filename_hint: str = "document.pdf",
    ) -> AdapterResult:
        if output_format != "docx":
            raise ConversionError(
                f"Local adapter only supports docx, not {output_format!r}. "
                "Use Adobe or LibreOffice for xlsx/pptx.",
                adapter_name=self.name,
                retryable=False,  # user picked a format we can't do
            )

        t0 = time.time()
        try:
            from docx import Document
        except ImportError as exc:
            raise ConversionError(
                "python-docx not installed on the server",
                adapter_name=self.name,
                cause=exc,
                retryable=False,
            ) from exc

        try:
            src = fitz.open(stream=pdf_bytes, filetype="pdf")
        except Exception as exc:
            raise ConversionError(
                f"Could not read PDF: {exc}",
                adapter_name=self.name,
                cause=exc,
                retryable=False,  # bad input, not transient
            ) from exc

        try:
            page_count = len(src)
            if page_count == 0:
                raise ConversionError(
                    "PDF has no pages",
                    adapter_name=self.name,
                    retryable=False,
                )

            out_doc = Document()
            first_page = True
            for page in src:
                if not first_page:
                    out_doc.add_page_break()
                first_page = False
                page_dict = page.get_text("dict")
                for block in page_dict.get("blocks", []):
                    if block.get("type", 0) == 1:
                        continue  # image block — skip
                    block_lines = block.get("lines", [])
                    for line in block_lines:
                        spans = line.get("spans", [])
                        if not spans:
                            continue
                        text = "".join(s.get("text", "") for s in spans).strip()
                        if not text:
                            continue
                        sizes = [s.get("size", 0) for s in spans if s.get("size", 0) > 0]
                        max_size = max(sizes) if sizes else 12
                        page_sizes = [
                            s.get("size", 0)
                            for b in page_dict.get("blocks", [])
                            for l in b.get("lines", [])
                            for s in l.get("spans", [])
                            if s.get("size", 0) > 0
                        ]
                        page_median = (
                            sorted(page_sizes)[len(page_sizes) // 2]
                            if page_sizes
                            else 12
                        )
                        if max_size >= page_median * 1.5:
                            out_doc.add_heading(text, level=1)
                        elif max_size >= page_median * 1.25:
                            out_doc.add_heading(text, level=2)
                        else:
                            out_doc.add_paragraph(text)
                    # Simple table detection
                    if _looks_like_table(block_lines):
                        rows = _extract_table_rows(block_lines)
                        if rows and len(rows) > 1:
                            ncols = max(len(r) for r in rows)
                            table = out_doc.add_table(rows=len(rows), cols=ncols)
                            for r_idx, row in enumerate(rows):
                                for c_idx, cell_text in enumerate(row):
                                    if c_idx < ncols:
                                        table.cell(r_idx, c_idx).text = cell_text

            out_buf = io.BytesIO()
            out_doc.save(out_buf)
            out_bytes = out_buf.getvalue()
        except ConversionError:
            raise
        except Exception as exc:
            raise ConversionError(
                f"Local conversion failed: {exc}",
                adapter_name=self.name,
                cause=exc,
                retryable=True,
            ) from exc
        finally:
            src.close()

        elapsed_ms = int((time.time() - t0) * 1000)
        return AdapterResult(
            bytes=out_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_extension="docx",
            adapter_name=self.name,
            elapsed_ms=elapsed_ms,
            pages_converted=page_count,
            cost_usd=0.0,
        )


def _looks_like_table(lines: list[dict[str, Any]]) -> bool:
    if len(lines) < 2:
        return False
    line_xs: list[set[float]] = []
    for line in lines:
        spans = line.get("spans", [])
        if not spans:
            line_xs.append(set())
            continue
        xs = {round(s.get("bbox", [0, 0, 0, 0])[0]) for s in spans}
        line_xs.append(xs)
    if not all(line_xs):
        return False
    counts = [len(xs) for xs in line_xs]
    if max(counts) < 2:
        return False
    if max(counts) - min(counts) > 1:
        return False
    return True


def _extract_table_rows(lines: list[dict[str, Any]]) -> list[list[str]]:
    all_xs: list[float] = []
    for line in lines:
        for s in line.get("spans", []):
            all_xs.append(s.get("bbox", [0, 0, 0, 0])[0])
    if not all_xs:
        return []
    sorted_xs = sorted(set(round(x) for x in all_xs))
    clusters: list[list[float]] = []
    for x in sorted_xs:
        if clusters and x - clusters[-1][-1] < 8:
            clusters[-1].append(x)
        else:
            clusters.append([x])
    col_centers = [sum(c) / len(c) for c in clusters]

    def col_for(x: float) -> int:
        best, best_dist = 0, abs(x - col_centers[0])
        for i, c in enumerate(col_centers):
            d = abs(x - c)
            if d < best_dist:
                best, best_dist = i, d
        return best

    rows: list[list[str]] = []
    for line in lines:
        cells: dict[int, list[str]] = {}
        for s in line.get("spans", []):
            x = s.get("bbox", [0, 0, 0, 0])[0]
            cells.setdefault(col_for(x), []).append(s.get("text", ""))
        if not cells:
            continue
        ncols = len(col_centers)
        row = ["".join(cells.get(i, [])) for i in range(ncols)]
        rows.append(row)
    return rows
