"""OCR (scanned PDF → searchable PDF or plain text) and PDF → Word endpoints.

OCR uses Tesseract via pytesseract. The OCR'd text is rendered as an
invisible text layer over the original scanned image, producing a
"searchable PDF" that looks identical but lets users Ctrl-F and
select/copy text.

PDF → Word uses python-docx. The conversion is best-effort: text,
headings, paragraphs, and basic tables. Multi-column layouts, math,
and complex tables will NOT round-trip faithfully. This is clearly
disclosed in the UI and the API response.
"""

from __future__ import annotations

import io
import logging
import time
from typing import Annotated, Literal

import fitz  # PyMuPDF
import pytesseract
from fastapi import APIRouter, File, Form, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from PIL import Image
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

router = APIRouter()

# Same 50 MB cap as the rest of the PDF router. OCR is CPU-bound and
# synchronous; larger files should go through the async queue.
MAX_SYNC_SIZE = 50 * 1024 * 1024


# ─── OCR ────────────────────────────────────────────────────────
class OcrTextResponse(BaseModel):
    text: str = Field(..., description="OCR'd plain text, one block per page")
    pages: int = Field(..., description="Number of pages processed")
    chars: int = Field(..., description="Total character count")
    elapsed_ms: int = Field(..., description="Time taken in milliseconds")


def _ocr_page(page: fitz.Page, lang: str = "eng", dpi: int = 300) -> str:
    """Render a page to a high-DPI image, OCR it, return plain text."""
    # Higher DPI = better OCR accuracy but slower. 300 is the standard
    # archival DPI. For very large documents 200 is a reasonable tradeoff.
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix, alpha=False)
    try:
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    finally:
        pix = None  # free C-level memory
    return pytesseract.image_to_string(img, lang=lang)


@router.post(
    "/ocr",
    response_model=OcrTextResponse,
    summary="OCR a scanned PDF and return plain text (≤ 50 MB)",
)
async def ocr_text(
    file: Annotated[UploadFile, File(description="Scanned PDF")],
    lang: Annotated[
        str,
        Form(description="Tesseract language code (default 'eng'). Multiple langs: 'eng+fra'."),
    ] = "eng",
    dpi: Annotated[
        int,
        Form(description="Render DPI for OCR (150-600, default 300)"),
    ] = 300,
) -> OcrTextResponse:
    """Run OCR on every page of a PDF and return the extracted text.

    This is the text-only endpoint. Use `/ocr-download` if you need a
    searchable PDF (PDF with an invisible text layer over the scans).
    """
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF.")
    if dpi < 150 or dpi > 600:
        raise HTTPException(400, "DPI must be between 150 and 600.")
    blob = await file.read()
    if len(blob) > MAX_SYNC_SIZE:
        raise HTTPException(413, f"File exceeds {MAX_SYNC_SIZE // (1024 * 1024)} MB limit.")
    if len(blob) == 0:
        raise HTTPException(400, "Empty file.")

    try:
        src = fitz.open(stream=blob, filetype="pdf")
    except Exception as exc:
        raise HTTPException(400, f"Could not read PDF: {exc}") from exc

    t0 = time.time()
    try:
        page_count = len(src)
        if page_count == 0:
            raise HTTPException(400, "PDF has no pages.")
        # OCR each page
        texts: list[str] = []
        for i, page in enumerate(src):
            try:
                text = _ocr_page(page, lang=lang, dpi=dpi)
            except pytesseract.TesseractNotFoundError as exc:
                raise HTTPException(
                    500,
                    "Tesseract binary not found on the server. "
                    "The OCR feature is misconfigured.",
                ) from exc
            except Exception as exc:
                logger.warning("OCR failed on page %d: %s", i + 1, exc)
                text = ""
            texts.append(text)
        full_text = "\n\n".join(texts)
    finally:
        src.close()
    elapsed_ms = int((time.time() - t0) * 1000)

    return OcrTextResponse(
        text=full_text,
        pages=page_count,
        chars=len(full_text),
        elapsed_ms=elapsed_ms,
    )


@router.post(
    "/ocr-download",
    response_class=StreamingResponse,
    summary="OCR a scanned PDF and return a searchable PDF (≤ 50 MB)",
    responses={
        200: {
            "description": "Searchable PDF (invisible text layer over the scans)",
            "content": {"application/pdf": {}},
        },
        400: {"description": "Invalid PDF or Tesseract error"},
        413: {"description": "File exceeds 50 MB cap"},
    },
)
async def ocr_pdf_download(
    file: Annotated[UploadFile, File(description="Scanned PDF")],
    lang: Annotated[
        str,
        Form(description="Tesseract language code (default 'eng'). Multiple: 'eng+fra'."),
    ] = "eng",
    dpi: Annotated[
        int,
        Form(description="Render DPI for OCR (150-600, default 300)"),
    ] = 300,
) -> StreamingResponse:
    """OCR a scanned PDF and return a searchable PDF.

    The output PDF looks identical to the input (we keep the original
    rendered page) but with an invisible text layer on top. This makes
    the document searchable (Ctrl+F) and lets users select/copy text.

    Trade-off: OCR is imperfect. Scanned documents with poor image
    quality, handwriting, complex layouts, or unusual fonts will
    produce noisy text. The visual appearance is unchanged, so the
    original is preserved — only the text layer is added.
    """
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF.")
    if dpi < 150 or dpi > 600:
        raise HTTPException(400, "DPI must be between 150 and 600.")
    blob = await file.read()
    if len(blob) > MAX_SYNC_SIZE:
        raise HTTPException(413, f"File exceeds {MAX_SYNC_SIZE // (1024 * 1024)} MB limit.")
    if len(blob) == 0:
        raise HTTPException(400, "Empty file.")

    try:
        src = fitz.open(stream=blob, filetype="pdf")
    except Exception as exc:
        raise HTTPException(400, f"Could not read PDF: {exc}") from exc

    t0 = time.time()
    try:
        page_count = len(src)
        if page_count == 0:
            raise HTTPException(400, "PDF has no pages.")
        zoom = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)

        for i, page in enumerate(src):
            try:
                # Render the page to a high-res image
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                try:
                    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                finally:
                    pix = None
                # OCR with bounding boxes via image_to_data, then map
                # boxes back to PDF coordinates and insert as invisible text.
                data = pytesseract.image_to_data(
                    img, lang=lang, output_type=pytesseract.Output.DICT
                )
                # PDF coordinates: we rendered at zoom = dpi/72, so divide back.
                inv_zoom = 72.0 / dpi
                words_inserted = 0
                for j in range(len(data["text"])):
                    word = data["text"][j].strip()
                    if not word:
                        continue
                    conf = data["conf"][j]
                    if isinstance(conf, str):
                        try:
                            conf = int(conf)
                        except ValueError:
                            conf = -1
                    if conf < 30:
                        # Skip low-confidence words — they pollute search
                        continue
                    x = data["left"][j] * inv_zoom
                    y = data["top"][j] * inv_zoom
                    w = data["width"][j] * inv_zoom
                    h = data["height"][j] * inv_zoom
                    # Insert as invisible text (white text, same color as bg
                    # is risky if the bg isn't pure white; use text rendering
                    # mode 3 which makes the glyph invisible regardless of color).
                    rect = fitz.Rect(x, y, x + w, y + h)
                    # Use a small fontsize that fits the bounding box.
                    fontsize = max(1.0, h * 0.85)
                    try:
                        page.insert_text(
                            (x, y + h * 0.85),  # baseline
                            word,
                            fontsize=fontsize,
                            color=(1, 1, 1),  # white — invisible on white scans
                            render_mode=3,  # text rendering mode 3 = invisible
                        )
                        words_inserted += 1
                    except Exception:
                        # Skip individual word insertion errors
                        pass
                logger.info("OCR page %d: %d words indexed", i + 1, words_inserted)
            except pytesseract.TesseractNotFoundError as exc:
                raise HTTPException(
                    500,
                    "Tesseract binary not found on the server.",
                ) from exc
            except Exception as exc:
                logger.warning("OCR failed on page %d: %s", i + 1, exc)
                # Continue — that page just won't have a text layer

        out_buf = io.BytesIO()
        src.save(out_buf, garbage=4, deflate=True, clean=True)
        out_bytes = out_buf.getvalue()
    finally:
        src.close()
    elapsed_ms = int((time.time() - t0) * 1000)

    base = (file.filename or "document.pdf").rsplit(".", 1)[0]
    out_name = f"{base}-searchable.pdf"

    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{out_name}"',
            "X-Pdf-Source-Pages": str(page_count),
            "X-Pdf-Size-Bytes": str(len(out_bytes)),
            "X-Ocr-Lang": lang,
            "X-Ocr-Dpi": str(dpi),
            "X-Ocr-Elapsed-Ms": str(elapsed_ms),
            "Cache-Control": "no-store",
        },
    )


# ─── PDF → Word (.docx) ────────────────────────────────────────
class ToWordResponse(BaseModel):
    pages: int = Field(..., description="Number of pages in the source PDF")
    paragraphs: int = Field(..., description="Number of paragraphs in the output .docx")
    tables: int = Field(..., description="Number of tables detected and converted")
    size_bytes: int = Field(..., description="Size of the .docx in bytes")
    accuracy_warning: str = Field(
        ...,
        description="Honest disclaimer about conversion accuracy",
    )


@router.post(
    "/to-word",
    response_model=ToWordResponse,
    summary="Convert a PDF to .docx (best-effort text + structure)",
)
async def pdf_to_word(
    file: Annotated[UploadFile, File(description="PDF to convert")],
) -> ToWordResponse:
    """Convert a PDF to a .docx file.

    **Honest accuracy note:** This is a *best-effort* conversion. We
    extract text, headings, and basic tables. The following will NOT
    round-trip faithfully:
      - Multi-column layouts
      - Complex tables (merged cells, nested tables)
      - Mathematical equations
      - Custom fonts and precise positioning
      - Vector graphics (we drop them)

    For complex documents, expect to do manual cleanup in Word. We
    do NOT claim to be a 1:1 conversion engine — the alternative
    is to lie to you, which we won't do.
    """
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF.")
    blob = await file.read()
    if len(blob) > MAX_SYNC_SIZE:
        raise HTTPException(413, f"File exceeds {MAX_SYNC_SIZE // (1024 * 1024)} MB limit.")
    if len(blob) == 0:
        raise HTTPException(400, "Empty file.")

    try:
        src = fitz.open(stream=blob, filetype="pdf")
    except Exception as exc:
        raise HTTPException(400, f"Could not read PDF: {exc}") from exc

    try:
        page_count = len(src)
        if page_count == 0:
            raise HTTPException(400, "PDF has no pages.")

        from docx import Document
        from docx.shared import Pt

        out_doc = Document()
        # Page break between pages (except before the first page)
        first_page = True
        total_paragraphs = 0
        total_tables = 0

        for page in src:
            if not first_page:
                out_doc.add_page_break()
            first_page = False

            # Extract structured text from the page
            # mode="dict" gives us per-block info: type (text/image), bbox, lines, spans
            page_dict = page.get_text("dict")

            for block in page_dict.get("blocks", []):
                block_type = block.get("type", 0)
                if block_type == 1:
                    # Image block — skip; we don't embed images in the .docx
                    continue
                # Text block
                block_lines = block.get("lines", [])
                if not block_lines:
                    continue
                # If a block has only one line and the line contains a long
                # run of single-char spans, treat it as a heading candidate.
                # Heuristic: if the average font size on the page is X and
                # this block's first line's largest font size is > 1.3X, it's
                # likely a heading.
                for line in block_lines:
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    if not text:
                        continue
                    # Use the largest font size in the line to pick a heading
                    # level. If it's notably larger than the page median, mark
                    # it as a heading.
                    sizes = [s.get("size", 0) for s in spans if s.get("size", 0) > 0]
                    max_size = max(sizes) if sizes else 12
                    page_sizes = [
                        s.get("size", 0)
                        for b in page_dict.get("blocks", [])
                        for l in b.get("lines", [])
                        for s in l.get("spans", [])
                        if s.get("size", 0) > 0
                    ]
                    page_median = sorted(page_sizes)[len(page_sizes) // 2] if page_sizes else 12

                    if max_size >= page_median * 1.5:
                        out_doc.add_heading(text, level=1)
                    elif max_size >= page_median * 1.25:
                        out_doc.add_heading(text, level=2)
                    else:
                        para = out_doc.add_paragraph(text)
                        # Set font size to the first span's size, falling
                        # back to 12. python-docx defaults to Calibri 11.
                        para.style = out_doc.styles["Normal"]
                    total_paragraphs += 1

                # After extracting all lines from this block, try to detect
                # a simple table: if the block's spans are arranged in a grid
                # (multiple lines, each with spans at roughly the same x
                # positions), treat the lines as table rows.
                if _looks_like_table(block_lines):
                    rows = _extract_table_rows(block_lines)
                    if rows and len(rows) > 1:
                        ncols = max(len(r) for r in rows)
                        table = out_doc.add_table(rows=len(rows), cols=ncols)
                        for r_idx, row in enumerate(rows):
                            for c_idx, cell_text in enumerate(row):
                                if c_idx < ncols:
                                    table.cell(r_idx, c_idx).text = cell_text
                        total_tables += 1

        out_buf = io.BytesIO()
        out_doc.save(out_buf)
        out_bytes = out_buf.getvalue()
    finally:
        src.close()

    return ToWordResponse(
        pages=page_count,
        paragraphs=total_paragraphs,
        tables=total_tables,
        size_bytes=len(out_bytes),
        accuracy_warning=(
            "Best-effort conversion. Text, headings, and basic tables are "
            "preserved. Multi-column layouts, math, complex tables, custom "
            "fonts, and precise positioning will NOT round-trip faithfully. "
            "Expect manual cleanup in Word for complex documents."
        ),
    )


@router.post(
    "/to-word-download",
    response_class=StreamingResponse,
    summary="Convert a PDF to .docx and stream the result back (≤ 50 MB)",
    responses={
        200: {
            "description": "Generated .docx file",
            "content": {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {}
            },
        },
    },
)
async def pdf_to_word_download(
    file: Annotated[UploadFile, File(description="PDF to convert")],
) -> StreamingResponse:
    """Same as /to-word but streams the .docx bytes directly. Use this
    from the web app; use /to-word for metadata-only calls.
    """
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF.")
    blob = await file.read()
    if len(blob) > MAX_SYNC_SIZE:
        raise HTTPException(413, f"File exceeds {MAX_SYNC_SIZE // (1024 * 1024)} MB limit.")
    if len(blob) == 0:
        raise HTTPException(400, "Empty file.")

    try:
        src = fitz.open(stream=blob, filetype="pdf")
    except Exception as exc:
        raise HTTPException(400, f"Could not read PDF: {exc}") from exc

    try:
        from docx import Document

        out_doc = Document()
        page_count = len(src)
        if page_count == 0:
            raise HTTPException(400, "PDF has no pages.")

        first_page = True
        for page in src:
            if not first_page:
                out_doc.add_page_break()
            first_page = False

            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type", 0) == 1:
                    continue
                block_lines = block.get("lines", [])
                for line in block_lines:
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    if not text:
                        continue
                    sizes = [s.get("size", 0) for s in spans if s.get("size", 0) > 0]
                    max_size = max(sizes) if sizes else 12
                    page_sizes = [
                        s.get("size", 0)
                        for b in page_dict.get("blocks", [])
                        for l in b.get("lines", [])
                        for s in l.get("spans", [])
                        if s.get("size", 0) > 0
                    ]
                    page_median = sorted(page_sizes)[len(page_sizes) // 2] if page_sizes else 12
                    if max_size >= page_median * 1.5:
                        out_doc.add_heading(text, level=1)
                    elif max_size >= page_median * 1.25:
                        out_doc.add_heading(text, level=2)
                    else:
                        out_doc.add_paragraph(text)
                if _looks_like_table(block_lines):
                    rows = _extract_table_rows(block_lines)
                    if rows and len(rows) > 1:
                        ncols = max(len(r) for r in rows)
                        table = out_doc.add_table(rows=len(rows), cols=ncols)
                        for r_idx, row in enumerate(rows):
                            for c_idx, cell_text in enumerate(row):
                                if c_idx < ncols:
                                    table.cell(r_idx, c_idx).text = cell_text

        out_buf = io.BytesIO()
        out_doc.save(out_buf)
        out_bytes = out_buf.getvalue()
    finally:
        src.close()

    base = (file.filename or "document.pdf").rsplit(".", 1)[0]
    out_name = f"{base}.docx"

    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{out_name}"',
            "X-Pdf-Source-Pages": str(page_count),
            "X-Docx-Size-Bytes": str(len(out_bytes)),
            "X-Conversion-Accuracy": "best-effort",
            "Cache-Control": "no-store",
        },
    )


def _looks_like_table(lines: list[dict]) -> bool:
    """Heuristic: a block looks like a table if there are ≥2 lines and
    each line has spans at consistent x positions (within a small tolerance)."""
    if len(lines) < 2:
        return False
    # Get the set of "column" x positions for each line
    line_xs: list[set[float]] = []
    for line in lines:
        spans = line.get("spans", [])
        if not spans:
            line_xs.append(set())
            continue
        xs = {round(s.get("bbox", [0, 0, 0, 0])[0]) for s in spans}
        line_xs.append(xs)
    if not all(line_xs):
        return False
    # All lines should have a similar number of "columns"
    counts = [len(xs) for xs in line_xs]
    if max(counts) < 2:
        return False
    if max(counts) - min(counts) > 1:
        return False
    return True


def _extract_table_rows(lines: list[dict]) -> list[list[str]]:
    """Group spans by their x-position clusters, return rows of cell text."""
    # Compute all x positions across all lines
    all_xs: list[float] = []
    for line in lines:
        for s in line.get("spans", []):
            all_xs.append(s.get("bbox", [0, 0, 0, 0])[0])
    if not all_xs:
        return []
    # Sort unique x positions, cluster within 5pt
    sorted_xs = sorted(set(round(x) for x in all_xs))
    clusters: list[list[float]] = []
    for x in sorted_xs:
        if clusters and x - clusters[-1][-1] < 8:
            clusters[-1].append(x)
        else:
            clusters.append([x])
    col_centers = [sum(c) / len(c) for c in clusters]

    def col_for(x: float) -> int:
        best, best_dist = 0, abs(x - col_centers[0])
        for i, c in enumerate(col_centers):
            d = abs(x - c)
            if d < best_dist:
                best, best_dist = i, d
        return best

    rows: list[list[str]] = []
    for line in lines:
        # Group spans by column
        cells: dict[int, list[str]] = {}
        for s in line.get("spans", []):
            x = s.get("bbox", [0, 0, 0, 0])[0]
            cells.setdefault(col_for(x), []).append(s.get("text", ""))
        if not cells:
            continue
        # Reorder by column index
        ncols = len(col_centers)
        row = ["".join(cells.get(i, [])) for i in range(ncols)]
        rows.append(row)
    return rows
